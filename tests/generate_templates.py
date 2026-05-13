#!/usr/bin/env python3
"""Generate 5 Russian .docx templates with different Jinja2 features."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color (hex string like 'D9E2F3')."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text):
    """Set cell text cleanly, avoiding duplicate text nodes that break docxtpl.

    python-docx's cell.text setter creates malformed XML where text appears
    both as a raw text node and inside <w:r><w:t>. docxtpl only parses
    <w:t> elements and gets confused by duplicates. We rebuild the cell
    from scratch with a single paragraph containing a single run.
    """
    tc = cell._tc
    # Clear all existing paragraphs
    for child in list(tc):
        if child.tag == qn('w:p'):
            tc.remove(child)

    # Create a new paragraph with a single run and text element
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)


def add_heading_paragraph(doc, text, level=1):
    """Add a heading-like paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def generate_template_01_variables():
    """Template 1: Simple variables (Договор/Contract)."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ДОГОВОР № {{ contract_number }}")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("г. {{ city }}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph()
    p.add_run("{{ contract_date }}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    doc.add_paragraph(
        "{{ company_name }}, именуемое в дальнейшем «Заказчик», в лице {{ signer_name }}, "
        "действующего на основании {{ signer_basis }}, с одной стороны, и ИП {{ contractor_name }}, "
        "именуемый в дальнейшем «Исполнитель», с другой стороны, заключили настоящий договор о нижеследующем:"
    )

    add_heading_paragraph(doc, "1. Предмет договора", 2)
    doc.add_paragraph(
        "1.1. Исполнитель обязуется оказать Заказчику услуги по {{ service_description }}, "
        "а Заказчик обязуется принять и оплатить эти услуги."
    )

    add_heading_paragraph(doc, "2. Стоимость услуг", 2)
    doc.add_paragraph(
        "2.1. Общая стоимость услуг составляет {{ total_amount }} рублей."
    )
    doc.add_paragraph(
        "2.2. Срок оказания услуг: {{ deadline }}."
    )

    add_heading_paragraph(doc, "3. Реквизиты сторон", 2)
    doc.add_paragraph("Заказчик: {{ company_name }}")
    doc.add_paragraph("ИНН: {{ company_inn }}")
    doc.add_paragraph("Адрес: {{ company_address }}")
    doc.add_paragraph()
    doc.add_paragraph("Исполнитель: {{ contractor_name }}")
    doc.add_paragraph("ИНН: {{ contractor_inn }}")
    doc.add_paragraph("Адрес: {{ contractor_address }}")

    doc.add_paragraph()
    doc.add_paragraph("Подпись Заказчика: _________________ / {{ signer_name }}")
    doc.add_paragraph("Подпись Исполнителя: _________________ / {{ contractor_name }}")

    doc.save("tests/input/template_01_variables.docx")


def generate_template_02_loops():
    """Template 2: Loops (Коммерческое предложение/Quotation with items)."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph("Уважаемый(ая) {{ client_name }}!")
    doc.add_paragraph()
    doc.add_paragraph(
        "Благодарим Вас за интерес к нашей компании. Предлагаем Вам следующие товары и услуги:"
    )
    doc.add_paragraph()

    doc.add_paragraph("{% for item in items %}")
    doc.add_paragraph(f"  • {{ item.name }} — {{ item.price }} ₽ ({{ item.description }})")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph()
    doc.add_paragraph("Итого позиций: {{ items|length }}")

    doc.add_paragraph()
    doc.add_paragraph("Дополнительные услуги:")
    doc.add_paragraph("{% for service in additional_services %}")
    doc.add_paragraph("  - {{ service.title }}: {{ service.cost }} ₽")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph()
    doc.add_paragraph("С уважением,")
    doc.add_paragraph("{{ manager_name }}")
    doc.add_paragraph("{{ manager_phone }}")
    doc.add_paragraph("{{ manager_email }}")
    doc.add_paragraph("{{ company_name }}")

    doc.save("tests/input/template_02_loops.docx")


def generate_template_03_conditionals():
    """Template 3: Conditionals (Счёт-фактура с условиями)."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СЧЁТ-ФАКТУРА")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph("Покупатель: {{ buyer_name }}")
    doc.add_paragraph("ИНН: {{ buyer_inn }}")
    doc.add_paragraph("Дата: {{ invoice_date }}")
    doc.add_paragraph("Номер счёта: {{ invoice_number }}")

    doc.add_paragraph()
    doc.add_paragraph("{% if is_vip %}")
    p = doc.add_paragraph()
    run = p.add_run("★ ПОКУПАТЕЛЬ VIP-УРОВНЯ ★")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph()
    doc.add_paragraph("Товары:")
    doc.add_paragraph("{% for product in products %}")
    doc.add_paragraph("  {{ loop.index }}. {{ product.name }} — {{ product.price }} ₽ × {{ product.quantity }} шт.")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph()
    doc.add_paragraph("Сумма без скидки: {{ subtotal }} ₽")

    doc.add_paragraph("{% if discount_percent > 0 %}")
    doc.add_paragraph("Скидка: {{ discount_percent }}%")
    doc.add_paragraph("Сумма скидки: {{ discount_amount }} ₽")
    doc.add_paragraph("{% else %}")
    doc.add_paragraph("Скидка не применена")
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("ИТОГО К ОПЛАТЕ: {{ total }} ₽")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("{% if include_vat %}")
    doc.add_paragraph("В том числе НДС 20%: {{ vat_amount }} ₽")
    doc.add_paragraph("{% else %}")
    doc.add_paragraph("НДС не облагается")
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph()
    doc.add_paragraph("{% if urgency == 'срочно' %}")
    p = doc.add_paragraph()
    run = p.add_run("⚡ СРОЧНЫЙ ЗАКАЗ — обработка в течение 24 часов!")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    doc.add_paragraph("{% elif urgency == 'обычный' %}")
    doc.add_paragraph("Стандартная обработка — 3-5 рабочих дней")
    doc.add_paragraph("{% else %}")
    doc.add_paragraph("Обычная обработка заказа")
    doc.add_paragraph("{% endif %}")

    doc.save("tests/input/template_03_conditionals.docx")


def generate_template_04_table():
    """Template 4: Table with row loops (Накладная/Waybill with table)."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("НАКЛАДНАЯ")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph("Отправитель: {{ sender_name }}")
    doc.add_paragraph("Получатель: {{ receiver_name }}")
    doc.add_paragraph("Дата отгрузки: {{ shipment_date }}")
    doc.add_paragraph("Номер накладной: {{ waybill_number }}")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    headers = ["№", "Наименование товара", "Кол-во", "Примечание"]
    for i, text in enumerate(headers):
        set_cell_text(hdr_cells[i], text)
        set_cell_shading(hdr_cells[i], "D9E2F3")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Row 1: loop start (this row is removed by docxtpl)
    row = table.add_row().cells
    set_cell_text(row[0], "{%tr for product in products %}")
    set_cell_text(row[1], "")
    set_cell_text(row[2], "")
    set_cell_text(row[3], "")

    # Row 2: template data row (this row gets duplicated for each item)
    data_row = table.add_row().cells
    set_cell_text(data_row[0], "{{ loop.index }}")
    set_cell_text(data_row[1], "{{ product.name }}")
    set_cell_text(data_row[2], "{{ product.quantity }}")
    set_cell_text(data_row[3], "{{ product.note }}")

    # Row 3: loop end (this row is removed by docxtpl)
    end_row = table.add_row().cells
    set_cell_text(end_row[0], "{%tr endfor %}")
    set_cell_text(end_row[1], "")
    set_cell_text(end_row[2], "")
    set_cell_text(end_row[3], "")

    doc.add_paragraph()
    doc.add_paragraph("Всего наименований: {{ products|length }}")
    doc.add_paragraph("Общее количество: {{ total_quantity }} ед.")

    doc.add_paragraph()
    doc.add_paragraph("Отгрузил: {{ shipper_name }} _________________")
    doc.add_paragraph("Получил: {{ receiver_signatory }} _________________")

    doc.save("tests/input/template_04_table.docx")


def generate_template_05_complex():
    """Template 5: Complex combined (Акт выполненных работ/Work completion certificate)."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АКТ ВЫПОЛНЕННЫХ РАБОТ")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("№ {{ act_number }} от {{ act_date }}")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph(
        "Мы, нижеподписавшиеся, представитель Заказчика {{ customer_name }} — "
        "{{ customer_representative }}, с одной стороны, и представитель Исполнителя "
        "{{ performer_name }} — {{ performer_representative }}, с другой стороны, "
        "составили настоящий акт о том, что Исполнителем выполнены следующие работы:"
    )

    doc.add_paragraph()

    # Table with row loops
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    headers = ["№ п/п", "Наименование работ", "Ед. изм.", "Кол-во", "Стоимость, ₽"]
    for i, text in enumerate(headers):
        set_cell_text(hdr_cells[i], text)
        set_cell_shading(hdr_cells[i], "D9E2F3")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Row 1: loop start (removed by docxtpl)
    row = table.add_row().cells
    set_cell_text(row[0], "{%tr for work in works %}")
    set_cell_text(row[1], "")
    set_cell_text(row[2], "")
    set_cell_text(row[3], "")
    set_cell_text(row[4], "")

    # Row 2: template data row (duplicated for each work)
    data_row = table.add_row().cells
    set_cell_text(data_row[0], "{{ loop.index }}")
    set_cell_text(data_row[1], "{{ work.name }}")
    set_cell_text(data_row[2], "{{ work.unit }}")
    set_cell_text(data_row[3], "{{ work.quantity }}")
    set_cell_text(data_row[4], "{{ work.cost }}")

    # Row 3: loop end (removed by docxtpl)
    end_row = table.add_row().cells
    set_cell_text(end_row[0], "{%tr endfor %}")
    set_cell_text(end_row[1], "")
    set_cell_text(end_row[2], "")
    set_cell_text(end_row[3], "")
    set_cell_text(end_row[4], "")

    doc.add_paragraph()

    # Totals with conditionals
    doc.add_paragraph("{% if material_costs > 0 %}")
    doc.add_paragraph("Стоимость материалов: {{ material_costs }} ₽")
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph("{% if has_transport %}")
    doc.add_paragraph("Транспортные расходы: {{ transport_cost }} ₽")
    doc.add_paragraph("{% endif %}")

    p = doc.add_paragraph()
    run = p.add_run("ИТОГО: {{ total_cost }} ₽")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("{% if total_cost > 100000 %}")
    p = doc.add_paragraph()
    run = p.add_run("Внимание: крупная сделка — требуется дополнительное согласование!")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph()
    doc.add_paragraph("Работы выполнены в полном объёме, качественно и в срок.")
    doc.add_paragraph("Заказчик претензий по объёму, качеству и срокам выполнения работ не имеет.")

    doc.add_paragraph()
    doc.add_paragraph("Дополнительные примечания:")
    doc.add_paragraph("{% for note in notes %}")
    doc.add_paragraph("  {{ loop.index }}. {{ note }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("От Исполнителя:                                От Заказчика:")
    doc.add_paragraph("{{ performer_representative }}                 {{ customer_representative }}")
    doc.add_paragraph("_________________                              _________________")
    doc.add_paragraph("М.П.                                           М.П.")

    doc.save("tests/input/template_05_complex.docx")


if __name__ == "__main__":
    generate_template_01_variables()
    generate_template_02_loops()
    generate_template_03_conditionals()
    generate_template_04_table()
    generate_template_05_complex()
    print("All 5 templates generated in tests/input/")
